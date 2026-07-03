from __future__ import annotations

import io
import os
import tempfile

import fitz
from docx import Document

from app.core.storage import StorageError
from app.infra.storage.memory_storage import MemoryStorageAdapter
from app.modules.assistant.skills.menu_import.document_loader import (
    DOCX_MIME,
    PDF_MIME,
    load_docx_text,
    load_menu_source_from_storage,
    load_pdf_pages,
)


def _minimal_pdf_bytes() -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Tacos al pastor $85")
    data = doc.tobytes()
    doc.close()
    return data


def _minimal_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Entradas")
    document.add_paragraph("Guacamole — $95")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_load_pdf_pages_renders_png_pages():
    pdf_bytes = _minimal_pdf_bytes()
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(pdf_bytes)
        tmp_path = tmp.name

    try:
        pages = load_pdf_pages(tmp_path)
    finally:
        os.unlink(tmp_path)

    assert len(pages) == 1
    png_bytes, media_type = pages[0]
    assert media_type == "image/png"
    assert png_bytes.startswith(b"\x89PNG")


def test_load_docx_text_extracts_paragraphs():
    docx_bytes = _minimal_docx_bytes()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    try:
        text = load_docx_text(tmp_path)
    finally:
        os.unlink(tmp_path)

    assert "Entradas" in text
    assert "Guacamole" in text


def test_load_menu_source_from_storage_pdf():
    storage = MemoryStorageAdapter()
    path = "restaurants/demo/import/menu_source/menu.pdf"
    storage.upload(path, _minimal_pdf_bytes(), PDF_MIME)

    payload = load_menu_source_from_storage(path, PDF_MIME, storage=storage)

    assert payload.text is None
    assert len(payload.pages) == 1
    assert payload.pages[0].media_type == "image/png"
    assert payload.pages[0].image_bytes.startswith(b"\x89PNG")


def test_load_menu_source_from_storage_docx():
    storage = MemoryStorageAdapter()
    path = "restaurants/demo/import/menu_source/menu.docx"
    storage.upload(path, _minimal_docx_bytes(), DOCX_MIME)

    payload = load_menu_source_from_storage(path, DOCX_MIME, storage=storage)

    assert payload.pages == []
    assert payload.text is not None
    assert "Guacamole" in payload.text


def test_load_menu_source_from_storage_missing_raises():
    storage = MemoryStorageAdapter()
    try:
        load_menu_source_from_storage("missing.pdf", PDF_MIME, storage=storage)
        assert False, "expected StorageError"
    except StorageError:
        pass
