"""Load menu source documents from storage for vision OCR or text extraction."""

from __future__ import annotations

import io
import tempfile
from dataclasses import dataclass, field

import fitz
from docx import Document

from app.core.storage import StorageError, StoragePort
from app.infra.storage.factory import build_storage

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_RENDER_DPI = 150
_RENDER_ZOOM = _RENDER_DPI / 72.0


@dataclass(frozen=True)
class VisionPage:
    image_bytes: bytes
    media_type: str


@dataclass
class MenuSourcePayload:
    pages: list[VisionPage] = field(default_factory=list)
    text: str | None = None


def load_pdf_pages(path: str) -> list[tuple[bytes, str]]:
    """Render each PDF page to PNG bytes at 150 DPI."""
    pages: list[tuple[bytes, str]] = []
    matrix = fitz.Matrix(_RENDER_ZOOM, _RENDER_ZOOM)
    with fitz.open(path) as doc:
        for page in doc:
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            pages.append((pixmap.tobytes("png"), "image/png"))
    return pages


def _load_pdf_pages_from_bytes(data: bytes) -> list[tuple[bytes, str]]:
    pages: list[tuple[bytes, str]] = []
    matrix = fitz.Matrix(_RENDER_ZOOM, _RENDER_ZOOM)
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            pages.append((pixmap.tobytes("png"), "image/png"))
    return pages


def load_docx_text(path: str) -> str:
    """Extract paragraph and table text from a DOCX file."""
    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _load_docx_text_from_bytes(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
        tmp.write(data)
        tmp.flush()
        return load_docx_text(tmp.name)


def load_menu_source_from_storage(
    storage_path: str,
    mime_type: str,
    *,
    storage: StoragePort | None = None,
) -> MenuSourcePayload:
    """Download a menu source from storage and prepare pages or text for extraction."""
    store = storage or build_storage()
    try:
        data = store.read(storage_path)
    except StorageError as exc:
        raise StorageError(f"Could not read menu source at {storage_path}") from exc

    mime = mime_type.strip().lower()
    if mime == PDF_MIME:
        page_tuples = _load_pdf_pages_from_bytes(data)
        return MenuSourcePayload(
            pages=[VisionPage(image_bytes=png, media_type=media) for png, media in page_tuples],
            text=None,
        )
    if mime == DOCX_MIME:
        return MenuSourcePayload(pages=[], text=_load_docx_text_from_bytes(data))
    if mime.startswith("image/"):
        return MenuSourcePayload(
            pages=[VisionPage(image_bytes=data, media_type=mime)],
            text=None,
        )
    raise ValueError(f"Unsupported menu source mime type: {mime_type}")


def write_bytes_to_temp_file(data: bytes, suffix: str) -> str:
    """Persist bytes to a temp file for path-based loaders (tests)."""
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(data)
        tmp.flush()
        return tmp.name
